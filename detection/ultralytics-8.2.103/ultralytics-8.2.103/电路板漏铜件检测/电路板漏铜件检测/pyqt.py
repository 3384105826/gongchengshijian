import sys
import os
import time
import cv2
import numpy as np
import pymysql
from PyQt5.QtWidgets import (QApplication, QWidget, QLabel, QVBoxLayout, QHBoxLayout,
                             QPushButton, QListWidget, QComboBox, QSpinBox, QProgressBar,
                             QFileDialog, QLineEdit, QMessageBox, QTableWidget, QTableWidgetItem,
                             QTabWidget, QGridLayout, QGroupBox, QSplitter)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5 import QtGui
from ultralytics import YOLO  # 导入YOLOv8库

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 导入matplotlib用于数据可视化
import matplotlib
matplotlib.use('Qt5Agg')
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 导入窗口分割检测相关函数
from 窗口分割检测 import sliding_window_detection, merge_detections, apply_segmentation_to_image

# ------------------- 数据库连接类 -------------------
class DatabaseManager:
    def __init__(self):
        self.conn = None
        self.current_user_id = None
    
    def connect(self):
        """连接到MySQL数据库"""
        try:
            # 尝试连接到MySQL服务器
            print("正在连接到MySQL数据库...")
            self.conn = pymysql.connect(
                host="localhost",
                port=3306,
                user="root",
                password="ojj060113A",
                charset="utf8mb4",
                cursorclass=pymysql.cursors.DictCursor
            )
            print("成功连接到MySQL服务器")
            
            # 检查数据库是否存在
            with self.conn.cursor() as cursor:
                cursor.execute("SHOW DATABASES LIKE 'cv_results'")
                result = cursor.fetchone()
                
                if not result:
                    print("数据库cv_results不存在，正在创建...")
                    cursor.execute("CREATE DATABASE cv_results CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci")
                    print("数据库cv_results创建成功")
            
            # 选择数据库
            self.conn.select_db("cv_results")
            print("成功选择数据库cv_results")
            
            # 创建必要的表
            if not self._create_tables():
                return False, "创建表失败"
            
            return True, "连接成功"
            
        except pymysql.MySQLError as err:
            return False, f"数据库连接失败：{err}\n请检查：1. MySQL服务是否运行 2. 用户名密码是否正确 3. 网络连接是否正常"
        except Exception as e:
            return False, f"未知错误：{e}"
    
    def _create_tables(self):
        """创建必要的表"""
        try:
            with self.conn.cursor() as cursor:
                # 创建用户表
                create_user_table = """
                CREATE TABLE IF NOT EXISTS users (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    username VARCHAR(50) NOT NULL UNIQUE,
                    password VARCHAR(255) NOT NULL,
                    email VARCHAR(100),
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
                cursor.execute(create_user_table)
                
                # 创建检测结果表（添加user_id字段）
                create_result_table = """
                CREATE TABLE IF NOT EXISTS detection_results (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    user_id INT NOT NULL,
                    image_name VARCHAR(255) NOT NULL,
                    total_defects INT NOT NULL,
                    detection_time DATETIME DEFAULT CURRENT_TIMESTAMP,
                    window_size VARCHAR(20),
                    overlap_ratio FLOAT,
                    output_path VARCHAR(500),
                    FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
                cursor.execute(create_result_table)
                
                # 创建缺陷详情表
                create_detail_table = """
                CREATE TABLE IF NOT EXISTS defect_details (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    result_id INT,
                    defect_index INT,
                    confidence FLOAT,
                    class_name VARCHAR(100),
                    bbox_x1 FLOAT,
                    bbox_y1 FLOAT,
                    bbox_x2 FLOAT,
                    bbox_y2 FLOAT,
                    center_x FLOAT,
                    center_y FLOAT,
                    FOREIGN KEY (result_id) REFERENCES detection_results(id) ON DELETE CASCADE
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
                """
                cursor.execute(create_detail_table)
                
            self.conn.commit()
            return True
        except Exception as e:
            print(f"创建表失败：{e}")
            return False
    
    def register_user(self, username, password, email=""):
        """用户注册"""
        if not self.conn or not self.conn.open:
            return False, "数据库未连接"
        
        try:
            with self.conn.cursor() as cursor:
                # 检查用户名是否已存在
                cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                if cursor.fetchone():
                    return False, "用户名已存在"
                
                # 插入新用户
                insert_sql = "INSERT INTO users (username, password, email) VALUES (%s, %s, %s)"
                cursor.execute(insert_sql, (username, password, email))
            
            self.conn.commit()
            return True, "注册成功"
        except Exception as e:
            self.conn.rollback()
            return False, f"注册失败：{e}"
    
    def verify_user(self, username, password):
        """验证用户登录"""
        if not self.conn or not self.conn.open:
            return None, "数据库未连接"
        
        try:
            with self.conn.cursor() as cursor:
                select_sql = "SELECT id, username FROM users WHERE username = %s AND password = %s"
                cursor.execute(select_sql, (username, password))
                result = cursor.fetchone()
                
                if result:
                    self.current_user_id = result['id']
                    return result, "登录成功"
                else:
                    return None, "用户名或密码错误"
        except Exception as e:
            return None, f"验证失败：{e}"
    
    def save_detection_result(self, image_name, total_defects, window_size, overlap_ratio, output_path, defects):
        """保存检测结果到数据库"""
        if not self.conn or not self.conn.open:
            return False, "数据库未连接"
        
        if not self.current_user_id:
            return False, "用户未登录"
        
        try:
            with self.conn.cursor() as cursor:
                # 插入检测结果（包含user_id）
                insert_result_sql = """
                INSERT INTO detection_results (user_id, image_name, total_defects, window_size, overlap_ratio, output_path)
                VALUES (%s, %s, %s, %s, %s, %s)
                """
                cursor.execute(insert_result_sql, (self.current_user_id, image_name, total_defects, window_size, overlap_ratio, output_path))
                
                # 获取插入的结果ID
                result_id = cursor.lastrowid
                
                # 插入缺陷详情
                insert_detail_sql = """
                INSERT INTO defect_details (result_id, defect_index, confidence, class_name, bbox_x1, bbox_y1, bbox_x2, bbox_y2, center_x, center_y)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                
                for i, defect in enumerate(defects):
                    box = defect['box']
                    cursor.execute(insert_detail_sql, (
                        result_id, i + 1, defect['conf'], defect['class'],
                        box[0], box[1], box[2], box[3],
                        (box[0] + box[2]) / 2, (box[1] + box[3]) / 2
                    ))
            
            self.conn.commit()
            return True, "保存成功"
        except Exception as e:
            self.conn.rollback()
            return False, f"保存失败：{e}"
    
    def get_all_results(self):
        """获取当前用户的所有检测结果"""
        if not self.conn or not self.conn.open:
            return None, "数据库未连接"
        
        if not self.current_user_id:
            return None, "用户未登录"
        
        try:
            with self.conn.cursor() as cursor:
                select_sql = "SELECT * FROM detection_results WHERE user_id = %s ORDER BY detection_time DESC"
                cursor.execute(select_sql, (self.current_user_id,))
                results = cursor.fetchall()
            
            return results, ""
        except Exception as e:
            return None, f"查询失败：{e}"
    
    def get_defect_details(self, result_id):
        """获取缺陷详情"""
        if not self.conn or not self.conn.open:
            return None, "数据库未连接"
        
        try:
            with self.conn.cursor() as cursor:
                select_sql = "SELECT * FROM defect_details WHERE result_id = %s"
                cursor.execute(select_sql, (result_id,))
                details = cursor.fetchall()
            
            return details, ""
        except Exception as e:
            return None, f"查询失败：{e}"
    
    def get_result_with_image(self, result_id):
        """获取检测结果和对应的图像"""
        if not self.conn or not self.conn.open:
            return None, "数据库未连接"
        
        try:
            with self.conn.cursor() as cursor:
                # 获取检测结果
                cursor.execute("SELECT * FROM detection_results WHERE id = %s", (result_id,))
                result = cursor.fetchone()
                
                if not result:
                    return None, "结果不存在"
                
                # 获取缺陷详情
                cursor.execute("SELECT * FROM defect_details WHERE result_id = %s", (result_id,))
                details = cursor.fetchall()
                
                # 读取输出图像
                output_path = result['output_path']
                image = None
                if output_path and os.path.exists(output_path):
                    image = cv2.imread(output_path)
                
                return {
                    'result': result,
                    'details': details,
                    'image': image
                }, ""
        except Exception as e:
            return None, f"查询失败：{e}"
    
    def delete_result(self, result_id):
        """删除检测结果及相关缺陷详情"""
        if not self.conn or not self.conn.open:
            return False, "数据库未连接"
        
        if not self.current_user_id:
            return False, "用户未登录"
        
        try:
            with self.conn.cursor() as cursor:
                # 检查结果是否属于当前用户
                cursor.execute("SELECT user_id FROM detection_results WHERE id = %s", (result_id,))
                result = cursor.fetchone()
                
                if not result:
                    return False, "结果不存在"
                
                if result['user_id'] != self.current_user_id:
                    return False, "无权删除此结果"
                
                # 开始事务
                self.conn.begin()
                
                # 删除相关的缺陷详情
                cursor.execute("DELETE FROM defect_details WHERE result_id = %s", (result_id,))
                
                # 删除检测结果
                cursor.execute("DELETE FROM detection_results WHERE id = %s", (result_id,))
                
                # 提交事务
                self.conn.commit()
                
                # 删除对应的输出文件
                cursor.execute("SELECT output_path FROM detection_results WHERE id = %s", (result_id,))
                output_path = cursor.fetchone()
                if output_path and output_path['output_path'] and os.path.exists(output_path['output_path']):
                    try:
                        os.remove(output_path['output_path'])
                    except:
                        pass
                
            return True, "删除成功"
        except Exception as e:
            self.conn.rollback()
            return False, f"删除失败：{e}"

    def close(self):
        """关闭数据库连接"""
        if self.conn and self.conn.open:
            self.conn.close()

# ------------------- 注册界面类 -------------------
class RegisterWindow(QWidget):
    register_success = pyqtSignal()
    
    def __init__(self, db_manager):
        super().__init__()
        self.db_manager = db_manager
        self.initUI()
    
    def initUI(self):
        self.setWindowTitle("用户注册")
        self.resize(400, 350)
        self.setMinimumSize(400, 350)
        
        self.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5;
                font-family: Arial, sans-serif;
                font-size: 14px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QLineEdit {
                padding: 8px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
                font-size: 14px;
            }
            QLabel {
                font-weight: bold;
                color: #333;
                font-size: 16px;
            }
            QWidget#registerWidget {
                background-color: white;
                border-radius: 10px;
                padding: 30px;
                box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            }
        """)
        
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)
        main_layout.setSpacing(20)
        
        register_widget = QWidget()
        register_widget.setObjectName("registerWidget")
        register_layout = QVBoxLayout(register_widget)
        register_layout.setSpacing(20)
        
        title_label = QLabel("用户注册")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("font-size: 20px; color: #4CAF50;")
        register_layout.addWidget(title_label)
        
        # 用户名
        user_layout = QHBoxLayout()
        user_layout.setSpacing(10)
        user_layout.addWidget(QLabel("用户名:"), 0)
        self.user_input = QLineEdit()
        self.user_input.setPlaceholderText("请输入用户名")
        user_layout.addWidget(self.user_input, 1)
        register_layout.addLayout(user_layout)
        
        # 密码
        pass_layout = QHBoxLayout()
        pass_layout.setSpacing(10)
        pass_layout.addWidget(QLabel("密码:"), 0)
        self.pass_input = QLineEdit()
        self.pass_input.setPlaceholderText("请输入密码")
        self.pass_input.setEchoMode(QLineEdit.Password)
        pass_layout.addWidget(self.pass_input, 1)
        register_layout.addLayout(pass_layout)
        
        # 确认密码
        confirm_layout = QHBoxLayout()
        confirm_layout.setSpacing(10)
        confirm_layout.addWidget(QLabel("确认密码:"), 0)
        self.confirm_input = QLineEdit()
        self.confirm_input.setPlaceholderText("请再次输入密码")
        self.confirm_input.setEchoMode(QLineEdit.Password)
        confirm_layout.addWidget(self.confirm_input, 1)
        register_layout.addLayout(confirm_layout)
        
        # 邮箱
        email_layout = QHBoxLayout()
        email_layout.setSpacing(10)
        email_layout.addWidget(QLabel("邮箱:"), 0)
        self.email_input = QLineEdit()
        self.email_input.setPlaceholderText("请输入邮箱（可选）")
        email_layout.addWidget(self.email_input, 1)
        register_layout.addLayout(email_layout)
        
        # 按钮
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        self.register_btn = QPushButton("注册")
        self.register_btn.setMinimumHeight(40)
        self.register_btn.clicked.connect(self.register)
        btn_layout.addWidget(self.register_btn)
        
        self.cancel_btn = QPushButton("取消")
        self.cancel_btn.setMinimumHeight(40)
        self.cancel_btn.setStyleSheet("background-color: #999;")
        self.cancel_btn.clicked.connect(self.close)
        btn_layout.addWidget(self.cancel_btn)
        
        register_layout.addLayout(btn_layout)
        
        main_layout.addWidget(register_widget)
        self.setLayout(main_layout)
    
    def register(self):
        """处理注册"""
        username = self.user_input.text().strip()
        password = self.pass_input.text().strip()
        confirm = self.confirm_input.text().strip()
        email = self.email_input.text().strip()
        
        if not username or not password:
            QMessageBox.warning(self, "注册失败", "用户名和密码不能为空！")
            return
        
        if len(username) < 3:
            QMessageBox.warning(self, "注册失败", "用户名长度不能少于3个字符！")
            return
        
        if len(password) < 6:
            QMessageBox.warning(self, "注册失败", "密码长度不能少于6个字符！")
            return
        
        if password != confirm:
            QMessageBox.warning(self, "注册失败", "两次输入的密码不一致！")
            return
        
        success, msg = self.db_manager.register_user(username, password, email)
        if success:
            QMessageBox.information(self, "注册成功", "用户注册成功！请登录。")
            self.register_success.emit()
            self.close()
        else:
            QMessageBox.warning(self, "注册失败", msg)


# ------------------- 登录界面类 -------------------
class LoginWindow(QWidget):
    login_success = pyqtSignal(str)

    def __init__(self, db_manager):
        super().__init__()
        self.db_manager = db_manager
        self.initUI()

    def initUI(self):
        self.setWindowTitle("用户登录")
        self.resize(400, 350)
        self.setMinimumSize(400, 350)

        self.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5;
                font-family: Arial, sans-serif;
                font-size: 14px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QLineEdit {
                padding: 8px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
                font-size: 14px;
            }
            QLabel {
                font-weight: bold;
                color: #333;
                font-size: 16px;
            }
            QWidget#loginWidget {
                background-color: white;
                border-radius: 10px;
                padding: 30px;
                box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            }
        """)

        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)
        main_layout.setSpacing(20)

        login_widget = QWidget()
        login_widget.setObjectName("loginWidget")
        login_layout = QVBoxLayout(login_widget)
        login_layout.setSpacing(20)

        title_label = QLabel("电路板漏铜检测系统")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("font-size: 20px; color: #4CAF50;")
        login_layout.addWidget(title_label)

        # 用户名
        user_layout = QHBoxLayout()
        user_layout.setSpacing(10)
        user_layout.addWidget(QLabel("用户名:"), 0)
        self.user_input = QLineEdit()
        self.user_input.setPlaceholderText("请输入用户名")
        user_layout.addWidget(self.user_input, 1)
        login_layout.addLayout(user_layout)

        # 密码
        pass_layout = QHBoxLayout()
        pass_layout.setSpacing(10)
        pass_layout.addWidget(QLabel("密码:"), 0)
        self.pass_input = QLineEdit()
        self.pass_input.setPlaceholderText("请输入密码")
        self.pass_input.setEchoMode(QLineEdit.Password)
        pass_layout.addWidget(self.pass_input, 1)
        login_layout.addLayout(pass_layout)

        # 按钮
        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(10)

        self.login_btn = QPushButton("登录")
        self.login_btn.setMinimumHeight(40)
        self.login_btn.clicked.connect(self.login)
        btn_layout.addWidget(self.login_btn)

        self.register_btn = QPushButton("注册新用户")
        self.register_btn.setMinimumHeight(40)
        self.register_btn.setStyleSheet("background-color: #2196F3;")
        self.register_btn.clicked.connect(self.show_register)
        btn_layout.addWidget(self.register_btn)

        login_layout.addLayout(btn_layout)

        main_layout.addWidget(login_widget)
        self.setLayout(main_layout)

    def login(self):
        """登录验证"""
        username = self.user_input.text().strip()
        password = self.pass_input.text().strip()

        if not username or not password:
            QMessageBox.warning(self, "登录失败", "用户名和密码不能为空！")
            return

        result, msg = self.db_manager.verify_user(username, password)
        if result:
            self.login_success.emit(result['username'])
            self.close()
        else:
            QMessageBox.warning(self, "登录失败", msg)

    def show_register(self):
        """显示注册窗口"""
        self.register_window = RegisterWindow(self.db_manager)
        self.register_window.register_success.connect(self.on_register_success)
        self.register_window.show()

    def on_register_success(self):
        """注册成功后"""
        QMessageBox.information(self, "提示", "注册成功！请使用新账号登录。")
    
    def login(self):
        """登录验证"""
        username = self.user_input.text().strip()
        password = self.pass_input.text().strip()
        
        if not username or not password:
            QMessageBox.warning(self, "登录失败", "用户名和密码不能为空！")
            return
        
        result, msg = self.db_manager.verify_user(username, password)
        if result:
            self.login_success.emit(result['username'])
            self.close()
        else:
            QMessageBox.warning(self, "登录失败", msg)
    
    def show_register(self):
        """显示注册窗口"""
        self.register_window = RegisterWindow(self.db_manager)
        self.register_window.register_success.connect(self.on_register_success)
        self.register_window.show()
    
    def on_register_success(self):
        """注册成功后"""
        QMessageBox.information(self, "提示", "注册成功！请使用新账号登录。")

# ------------------- 检测线程（异步处理，避免界面卡顿） -------------------
class YoloDetectThread(QThread):
    # 信号：更新进度、处理完成
    update_file_progress = pyqtSignal(int)
    update_batch_progress = pyqtSignal(int)
    process_done = pyqtSignal()
    file_done = pyqtSignal(str)

    def __init__(self, file_list, model_path, device, conf, output_dir, db_manager):
        super().__init__()
        self.file_list = file_list  # 待检测文件列表
        self.model_path = model_path  # 自定义模型路径（用户选择）
        self.device = device  # 设备（cpu/cuda）
        self.conf = conf / 100  # 置信度（转成0-1的浮点数）
        self.output_dir = output_dir  # 结果输出目录
        self.db_manager = db_manager  # 数据库管理器
        self.is_running = True  # 控制线程停止

    def run(self):
        try:
            # 加载YOLOv8模型（添加异常捕获，避免模型路径错误崩溃）
            model = YOLO(self.model_path)
        except Exception as e:
            QMessageBox.critical(None, "模型加载失败", f"错误原因：{str(e)}")
            self.process_done.emit()
            return

        total_files = len(self.file_list)
        for idx, file_path in enumerate(self.file_list):
            if not self.is_running:
                break  # 收到停止信号，中断处理

            # 1. 更新批处理进度
            batch_progress = int((idx + 1) / total_files * 100)
            self.update_batch_progress.emit(batch_progress)

            # 2. 单文件检测（使用滑动窗口检测算法）
            try:
                # 读取图像
                image = cv2.imread(file_path)
                if image is None:
                    QMessageBox.warning(None, "检测失败", f"无法读取图像：{os.path.basename(file_path)}")
                    self.update_file_progress.emit(0)
                    continue

                # 使用滑动窗口进行检测
                detections = sliding_window_detection(model, image, window_size=1280, conf_threshold=self.conf)

                # 合并重叠的检测结果
                merged_detections = merge_detections(detections)

                # 应用分割结果到图像
                result_image = apply_segmentation_to_image(image, merged_detections)

                # 保存结果图像
                os.makedirs(self.output_dir, exist_ok=True)
                output_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_segmented.bmp"
                output_path = os.path.join(self.output_dir, output_filename)
                cv2.imwrite(output_path, result_image)

                # 保存检测信息到文本文件
                info_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_detection_info.txt"
                info_path = os.path.join(self.output_dir, info_filename)

                with open(info_path, 'w', encoding='utf-8') as f:
                    f.write(f"图像: {os.path.basename(file_path)}\n")
                    f.write(f"总缺陷数: {len(merged_detections)}\n")
                    f.write(f"窗口大小: 1280*1280\n")
                    f.write(f"重叠比例: 0.25\n\n")

                    for j, det in enumerate(merged_detections):
                        box = det['box']
                        f.write(f"缺陷 {j + 1}:\n")
                        f.write(f"  置信度: {det['conf']:.3f}\n")
                        f.write(f"  类别: {det['class']}\n")
                        f.write(f"  边界框: ({box[0]:.1f}, {box[1]:.1f}, {box[2]:.1f}, {box[3]:.1f})\n")
                        f.write(f"  中心点: ({(box[0] + box[2]) / 2:.1f}, {(box[1] + box[3]) / 2:.1f})\n\n")

                # 推送保存的输出文件路径
                self.file_done.emit(output_path)
                
                # 保存到数据库
                if self.db_manager:
                    image_name = os.path.basename(file_path)
                    window_size = "1280*1280"
                    overlap_ratio = 0.25
                    
                    success, msg = self.db_manager.save_detection_result(
                        image_name=image_name,
                        total_defects=len(merged_detections),
                        window_size=window_size,
                        overlap_ratio=overlap_ratio,
                        output_path=output_path,
                        defects=merged_detections
                    )
                    if not success:
                        print(f"数据库保存失败: {msg}")

            except Exception as e:
                QMessageBox.warning(None, "检测失败", f"文件 {os.path.basename(file_path)} 检测出错：{str(e)}")
                self.update_file_progress.emit(0)
                continue

            # 3. 模拟单文件进度（实际可根据推理耗时调整）
            for i in range(101):
                if not self.is_running:
                    break
                self.update_file_progress.emit(i)
                time.sleep(0.005)  # 模拟处理耗时

        # 处理完成，发送信号
        self.process_done.emit()

    def stop(self):
        self.is_running = False


# ------------------- 主界面类 -------------------
class YoloV8Interface(QWidget):
    def __init__(self, db_manager, username):
        super().__init__()
        self.detect_thread = None  # 检测线程
        self.model_path = ""  # 初始化模型路径为空
        self.db_manager = db_manager  # 数据库管理器
        self.username = username  # 当前用户名
        self.current_result_id = None  # 当前选中的结果ID
        self.current_image = None  # 当前选中的图像
        self.initUI()
    
    def initUI(self):
        # 窗口设置
        self.setWindowTitle(f"电路板漏铜检测系统 - 当前用户: {self.username}")
        # 使用黄金分割率设置窗口大小 (1:1.618)
        height = 800
        width = int(height * 1.618)  # 黄金分割率
        self.resize(width, height)  # 调整窗口大小
        self.setMinimumSize(1000, 700)  # 设置最小窗口大小
        
        # 设置样式
        self.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5;
                font-family: Arial, sans-serif;
                font-size: 14px;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
            QLineEdit {
                padding: 6px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
            QListWidget {
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
            QProgressBar {
                border: 1px solid #ddd;
                border-radius: 4px;
                text-align: center;
                background-color: white;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                border-radius: 4px;
            }
            QComboBox {
                padding: 6px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
            QSpinBox {
                padding: 6px;
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
            QLabel {
                font-weight: bold;
                color: #333;
            }
            QTableWidget {
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
            QTableWidget::item {
                padding: 4px;
            }
            QTableWidget::header {
                background-color: #f0f0f0;
                font-weight: bold;
            }
            QScrollArea {
                border: 1px solid #ddd;
                border-radius: 4px;
                background-color: white;
            }
        """)

        # 用户信息栏
        user_bar_layout = QHBoxLayout()
        user_bar_layout.setSpacing(10)
        
        self.user_label = QLabel(f"当前用户: {self.username}")
        self.user_label.setStyleSheet("color: #4CAF50; font-size: 14px;")
        user_bar_layout.addWidget(self.user_label)
        
        user_bar_layout.addStretch()
        
        self.btn_logout = QPushButton("退出登录")
        self.btn_logout.setStyleSheet("background-color: #f44336;")
        self.btn_logout.clicked.connect(self.logout)
        user_bar_layout.addWidget(self.btn_logout)

        # 创建标签页
        self.tab_widget = QTabWidget()
        
        # 检测标签页
        detect_tab = QWidget()
        detect_layout = QVBoxLayout(detect_tab)
        detect_layout.setSpacing(20)
        detect_layout.setContentsMargins(20, 20, 20, 20)
        
        # 1. 顶部按钮栏（添加文件/文件夹、清空、打开输出）
        top_layout = QHBoxLayout()
        top_layout.setSpacing(10)
        
        self.btn_add_file = QPushButton("添加文件...")
        self.btn_add_file.clicked.connect(self.add_file)
        
        self.btn_add_folder = QPushButton("添加文件夹...")
        self.btn_add_folder.clicked.connect(self.add_folder)
        
        self.btn_clear = QPushButton("清空")
        self.btn_clear.clicked.connect(self.clear_list)
        
        self.btn_open_output = QPushButton("打开输出目录")
        self.btn_open_output.clicked.connect(self.open_output_dir)
        
        top_layout.addWidget(self.btn_add_file)
        top_layout.addWidget(self.btn_add_folder)
        top_layout.addStretch()
        top_layout.addWidget(self.btn_clear)
        top_layout.addWidget(self.btn_open_output)

        # 2. 文件列表区域（支持拖拽）
        file_list_group = QVBoxLayout()
        file_list_group.addWidget(QLabel("待检测文件："))
        self.file_list = QListWidget()
        self.file_list.setAcceptDrops(True)
        self.file_list.dragEnterEvent = self.drag_enter
        self.file_list.dragMoveEvent = self.drag_move
        self.file_list.dropEvent = self.drop_file
        self.file_list.setMinimumHeight(150)  # 调整高度
        file_list_group.addWidget(self.file_list)

        # 2.1 输出结果列表（显示检测后文件，双击打开）
        output_list_group = QVBoxLayout()
        output_list_group.addWidget(QLabel("输出结果："))
        self.output_list = QListWidget()
        self.output_list.itemDoubleClicked.connect(self.open_output_item)
        self.output_list.setMinimumHeight(100)  # 调整高度
        output_list_group.addWidget(self.output_list)

        # 3. 参数设置区域
        param_layout = QVBoxLayout()
        param_layout.setSpacing(15)
        param_layout.addWidget(QLabel("检测参数设置："))
        
        # 模型路径选择
        model_layout = QHBoxLayout()
        model_layout.setSpacing(10)
        model_layout.addWidget(QLabel("模型路径："), 0)
        self.le_model_path = QLineEdit()
        self.le_model_path.setPlaceholderText("请选择YOLOv11模型文件（.pt格式）")
        # 默认填充原路径（可选，方便用户直接使用）
        default_model_path = os.path.join(os.getcwd(), "model", "best.pt")
        if os.path.exists(default_model_path):
            self.le_model_path.setText(default_model_path)
            self.model_path = default_model_path
        self.le_model_path.setMinimumWidth(500)  # 调整宽度
        model_layout.addWidget(self.le_model_path, 1)
        self.btn_choose_model = QPushButton("选择模型...")
        self.btn_choose_model.clicked.connect(self.choose_model)
        model_layout.addWidget(self.btn_choose_model, 0)
        param_layout.addLayout(model_layout)

        # 设备选择（CPU/GPU）
        device_layout = QHBoxLayout()
        device_layout.setSpacing(10)
        device_layout.addWidget(QLabel("设备："), 0)
        self.cb_device = QComboBox()
        self.cb_device.addItems(["cpu", "cuda"])
        # 自动检测GPU，默认选cuda（如果可用）
        self.cb_device.setCurrentText("cuda" if self.check_cuda() else "cpu")
        device_layout.addWidget(self.cb_device, 1)
        param_layout.addLayout(device_layout)

        # 置信度阈值（0-100%）
        conf_layout = QHBoxLayout()
        conf_layout.setSpacing(10)
        conf_layout.addWidget(QLabel("置信度阈值："), 0)
        self.spin_conf = QSpinBox()
        self.spin_conf.setRange(0, 100)
        self.spin_conf.setValue(50)  # 默认50%
        conf_layout.addWidget(self.spin_conf, 1)
        conf_layout.addWidget(QLabel("%"), 0)
        param_layout.addLayout(conf_layout)

        # 输出目录
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)
        output_layout.addWidget(QLabel("输出目录："), 0)
        self.le_output = QLineEdit()
        self.le_output.setText(os.path.join(os.getcwd(), "yolo_output"))
        os.makedirs(self.le_output.text(), exist_ok=True)
        self.le_output.setMinimumWidth(500)  # 调整宽度
        output_layout.addWidget(self.le_output, 1)
        self.btn_browse = QPushButton("浏览...")
        self.btn_browse.clicked.connect(self.choose_output_dir)
        output_layout.addWidget(self.btn_browse, 0)
        param_layout.addLayout(output_layout)

        # 4. 进度条区域
        progress_layout = QVBoxLayout()
        progress_layout.setSpacing(10)
        progress_layout.addWidget(QLabel("检测进度："))
        
        self.pb_file = QProgressBar()
        self.pb_file.setFormat("当前文件进度：%p%")
        progress_layout.addWidget(self.pb_file)
        
        self.pb_batch = QProgressBar()
        self.pb_batch.setFormat("批处理进度：%p%")
        progress_layout.addWidget(self.pb_batch)

        # 5. 底部按钮（开始/停止）
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(30)  # 调整间距
        btn_layout.addStretch()
        
        self.btn_start = QPushButton("开始检测")
        self.btn_start.setMinimumSize(150, 45)  # 调整按钮大小
        self.btn_start.clicked.connect(self.start_detect)
        
        self.btn_stop = QPushButton("停止")
        self.btn_stop.setMinimumSize(150, 45)  # 调整按钮大小
        self.btn_stop.clicked.connect(self.stop_detect)
        self.btn_stop.setEnabled(False)
        
        btn_layout.addWidget(self.btn_start)
        btn_layout.addWidget(self.btn_stop)
        btn_layout.addStretch()

        # 检测标签页布局
        detect_layout.addLayout(top_layout)
        detect_layout.addLayout(file_list_group)
        detect_layout.addLayout(output_list_group)
        detect_layout.addLayout(param_layout)
        detect_layout.addLayout(progress_layout)
        detect_layout.addLayout(btn_layout)

        # 数据查看标签页
        data_tab = QWidget()
        data_layout = QVBoxLayout(data_tab)
        data_layout.setSpacing(20)
        data_layout.setContentsMargins(20, 20, 20, 20)
        
        # 数据查看标题
        data_layout.addWidget(QLabel("检测结果数据查看"))

        # 刷新按钮、删除按钮和导出按钮
        refresh_layout = QHBoxLayout()
        refresh_layout.addStretch()
        self.btn_refresh_data = QPushButton("刷新数据")
        self.btn_refresh_data.clicked.connect(self.refresh_data)
        refresh_layout.addWidget(self.btn_refresh_data)

        self.btn_delete_result = QPushButton("删除选中结果")
        self.btn_delete_result.setStyleSheet("background-color: #f44336;")
        self.btn_delete_result.clicked.connect(self.delete_selected_result)
        refresh_layout.addWidget(self.btn_delete_result)

        # 添加导出Word按钮
        self.btn_export_word = QPushButton("导出Word报告")
        self.btn_export_word.setStyleSheet("background-color: #2196F3;")
        self.btn_export_word.clicked.connect(self.export_to_word)
        refresh_layout.addWidget(self.btn_export_word)

        data_layout.addLayout(refresh_layout)
        
        # 检测结果和可视化区域
        content_layout = QHBoxLayout()
        content_layout.setSpacing(20)
        
        # 左侧：检测结果表格
        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel("检测结果列表："))
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(6)
        self.result_table.setHorizontalHeaderLabels(["ID", "图像名称", "缺陷数量", "检测时间", "窗口大小", "输出路径"])
        self.result_table.setMinimumHeight(200)
        self.result_table.horizontalHeader().setStretchLastSection(True)
        self.result_table.itemClicked.connect(self.on_result_item_clicked)
        self.result_table.setSelectionBehavior(QTableWidget.SelectRows)
        left_layout.addWidget(self.result_table)
        
        # 右侧：可视化区域
        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel("检测结果可视化："))
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setMinimumSize(400, 300)
        self.image_label.setStyleSheet("border: 1px solid #ddd; background-color: #f0f0f0;")
        self.image_label.setText("请选择一个检测结果查看可视化")
        right_layout.addWidget(self.image_label)
        
        content_layout.addLayout(left_layout, 3)
        content_layout.addLayout(right_layout, 2)
        data_layout.addLayout(content_layout)
        
        # 缺陷详情表格
        data_layout.addWidget(QLabel("缺陷详情："))
        self.detail_table = QTableWidget()
        self.detail_table.setColumnCount(6)
        self.detail_table.setHorizontalHeaderLabels(["缺陷序号", "置信度", "类别", "边界框", "中心点X", "中心点Y"])
        self.detail_table.setMinimumHeight(150)
        self.detail_table.horizontalHeader().setStretchLastSection(True)
        self.detail_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.detail_table.itemClicked.connect(self.on_detail_item_clicked)
        data_layout.addWidget(self.detail_table)

        # 数据可视化标签页
        visual_tab = QWidget()
        visual_layout = QVBoxLayout(visual_tab)
        visual_layout.setSpacing(20)
        visual_layout.setContentsMargins(20, 20, 20, 20)
        
        # 标题
        visual_layout.addWidget(QLabel("数据可视化分析"))
        
        # 刷新按钮
        visual_refresh_layout = QHBoxLayout()
        visual_refresh_layout.addStretch()
        self.btn_refresh_visual = QPushButton("刷新数据")
        self.btn_refresh_visual.clicked.connect(self.refresh_visualization)
        visual_refresh_layout.addWidget(self.btn_refresh_visual)
        visual_layout.addLayout(visual_refresh_layout)
        
        # 图表区域
        splitter = QSplitter(Qt.Vertical)
        
        # 第一行：缺陷数量统计
        defect_count_group = QGroupBox("缺陷数量统计")
        defect_count_layout = QHBoxLayout(defect_count_group)
        
        # 缺陷数量柱状图
        self.defect_count_canvas = FigureCanvas(Figure(figsize=(6, 4)))
        self.defect_count_ax = self.defect_count_canvas.figure.add_subplot(111)
        defect_count_layout.addWidget(self.defect_count_canvas)
        splitter.addWidget(defect_count_group)
        
        # 第二行：良品率和缺陷分布
        bottom_group = QGroupBox("综合分析")
        bottom_layout = QHBoxLayout(bottom_group)
        
        # 良品率饼图
        self.yield_pie_canvas = FigureCanvas(Figure(figsize=(4, 4)))
        self.yield_pie_ax = self.yield_pie_canvas.figure.add_subplot(111)
        bottom_layout.addWidget(self.yield_pie_canvas)
        
        # 缺陷位置分布热力图
        self.defect_dist_canvas = FigureCanvas(Figure(figsize=(6, 4)))
        self.defect_dist_ax = self.defect_dist_canvas.figure.add_subplot(111)
        bottom_layout.addWidget(self.defect_dist_canvas)
        splitter.addWidget(bottom_group)
        
        visual_layout.addWidget(splitter)
        
        # 统计信息
        stats_group = QGroupBox("统计信息")
        stats_layout = QGridLayout(stats_group)
        
        self.total_images_label = QLabel("总图像数: 0")
        self.total_defects_label = QLabel("总缺陷数: 0")
        self.average_defects_label = QLabel("平均缺陷数: 0")
        self.yield_rate_label = QLabel("良品率: 0%")
        self.pass_count_label = QLabel("合格品数: 0")
        self.fail_count_label = QLabel("不合格品数: 0")
        
        stats_layout.addWidget(self.total_images_label, 0, 0)
        stats_layout.addWidget(self.total_defects_label, 0, 1)
        stats_layout.addWidget(self.average_defects_label, 1, 0)
        stats_layout.addWidget(self.yield_rate_label, 1, 1)
        stats_layout.addWidget(self.pass_count_label, 2, 0)
        stats_layout.addWidget(self.fail_count_label, 2, 1)
        
        visual_layout.addWidget(stats_group)
        
        # 添加标签页
        self.tab_widget.addTab(detect_tab, "检测功能")
        self.tab_widget.addTab(data_tab, "数据查看")
        self.tab_widget.addTab(visual_tab, "数据可视化")

        # 整体布局
        main_layout = QVBoxLayout()
        main_layout.addLayout(user_bar_layout)
        main_layout.addWidget(self.tab_widget)
        
        self.setLayout(main_layout)
        
        # 初始化数据
        self.refresh_data()
        self.refresh_visualization()

    # ---------- 新增：模型路径选择方法 ----------
    def choose_model(self):
        """选择YOLOv8模型文件（.pt格式）"""
        model_file, _ = QFileDialog.getOpenFileName(
            self, "选择模型文件", "", "YOLO模型文件 (*.pt)"
        )
        if model_file:
            self.le_model_path.setText(model_file)
            self.model_path = model_file  # 更新模型路径

    # ---------- 辅助工具方法 ----------
    def check_cuda(self):
        """检查是否有可用GPU"""
        try:
            import torch
            return torch.cuda.is_available()
        except ImportError:
            return False

    def choose_output_dir(self):
        """选择输出目录"""
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if dir_path:
            self.le_output.setText(dir_path)

    def open_output_dir(self):
        """打开输出目录（Windows）"""
        dir_path = self.le_output.text()
        if os.path.exists(dir_path):
            os.startfile(dir_path)

    # ---------- 文件操作方法 ----------
    def add_file(self):
        """添加单个文件（图片/视频）"""
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", "", "支持格式 (*.jpg *.jpeg *.png *.bmp *.mp4 *.avi *.mov *.mkv)"
        )
        if files:
            self.file_list.addItems(files)

    def add_folder(self):
        """添加文件夹下的所有图片/视频"""
        folder = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder:
            supported_ext = (".jpg", ".jpeg", ".png", ".bmp", ".mp4", ".avi", ".mov", ".mkv")
            files = [
                os.path.join(folder, f)
                for f in os.listdir(folder)
                if f.lower().endswith(supported_ext)
            ]
            if files:
                self.file_list.addItems(files)

    def clear_list(self):
        """清空文件列表和进度条"""
        self.file_list.clear()
        self.output_list.clear()
        self.pb_file.setValue(0)
        self.pb_batch.setValue(0)

    # ---------- 拖拽文件方法 ----------
    def drag_enter(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def drag_move(self, event):
        event.acceptProposedAction()

    def drop_file(self, event):
        """拖拽文件到列表"""
        file_paths = [url.toLocalFile() for url in event.mimeData().urls()]
        self.file_list.addItems(file_paths)

    # ---------- 检测逻辑方法（修改模型路径获取） ----------
    def start_detect(self):
        """启动检测线程"""
        # 1. 检查必要参数
        if self.file_list.count() == 0:
            QMessageBox.warning(self, "提示", "请先添加要检测的文件！")
            return
        if not self.model_path or not os.path.exists(self.model_path):
            QMessageBox.warning(self, "提示", "请选择有效的模型文件（.pt格式）！")
            return

        # 2. 获取界面参数
        file_list = [self.file_list.item(i).text() for i in range(self.file_list.count())]
        device = self.cb_device.currentText()
        conf = self.spin_conf.value()
        output_dir = self.le_output.text()

        # 3. 启动检测线程
        self.detect_thread = YoloDetectThread(
            file_list=file_list,
            model_path=self.model_path,
            device=device,
            conf=conf,
            output_dir=output_dir,
            db_manager=self.db_manager
        )
        # 绑定信号与进度条更新
        self.detect_thread.update_file_progress.connect(self.pb_file.setValue)
        self.detect_thread.update_batch_progress.connect(self.pb_batch.setValue)
        self.detect_thread.process_done.connect(self.on_detect_done)
        self.detect_thread.file_done.connect(self.on_file_done)
        self.detect_thread.start()

        # 4. 按钮状态切换
        self.btn_start.setEnabled(False)
        self.btn_stop.setEnabled(True)

    def stop_detect(self):
        """停止检测线程"""
        if self.detect_thread:
            self.detect_thread.stop()
            self.detect_thread.wait()
            self.on_detect_done()

    def on_detect_done(self):
        """检测完成后的状态重置"""
        self.btn_start.setEnabled(True)
        self.btn_stop.setEnabled(False)
        QMessageBox.information(self, "提示", "检测完成！结果已保存到输出目录")

    def on_file_done(self, out_path: str):
        """单个文件检测完成后，显示输出路径"""
        self.output_list.addItem(out_path)

    def open_output_item(self, item):
        """双击打开某个输出文件"""
        path = item.text()
        if os.path.exists(path):
            try:
                os.startfile(path)
            except Exception:
                pass
    
    def refresh_data(self):
        """刷新检测结果数据"""
        if not self.db_manager:
            QMessageBox.warning(self, "提示", "数据库未连接")
            return
        
        # 获取所有检测结果
        results, msg = self.db_manager.get_all_results()
        if results is None:
            QMessageBox.warning(self, "错误", f"获取数据失败：{msg}")
            return
        
        # 清空表格
        self.result_table.setRowCount(0)
        
        # 填充表格
        for result in results:
            row = self.result_table.rowCount()
            self.result_table.insertRow(row)
            
            self.result_table.setItem(row, 0, QTableWidgetItem(str(result['id'])))
            self.result_table.setItem(row, 1, QTableWidgetItem(result['image_name']))
            self.result_table.setItem(row, 2, QTableWidgetItem(str(result['total_defects'])))
            self.result_table.setItem(row, 3, QTableWidgetItem(str(result['detection_time'])))
            self.result_table.setItem(row, 4, QTableWidgetItem(result['window_size']))
            self.result_table.setItem(row, 5, QTableWidgetItem(result['output_path']))
        
        # 清空详情表格
        self.detail_table.setRowCount(0)
    
    def on_result_item_clicked(self, item):
        """点击结果项，显示缺陷详情和可视化"""
        row = item.row()
        result_id = int(self.result_table.item(row, 0).text())
        self.current_result_id = result_id
        
        # 获取检测结果和图像
        data, msg = self.db_manager.get_result_with_image(result_id)
        if data is None:
            QMessageBox.warning(self, "错误", f"获取数据失败：{msg}")
            return
        
        # 显示图像可视化
        image = data['image']
        if image is not None:
            self.current_image = image
            # 将OpenCV图像转换为Qt图像
            rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            h, w, c = rgb_image.shape
            bytes_per_line = c * w
            qt_image = QtGui.QImage(rgb_image.data, w, h, bytes_per_line, QtGui.QImage.Format_RGB888)
            pixmap = QtGui.QPixmap.fromImage(qt_image)
            
            # 缩放图像以适应标签大小
            scaled_pixmap = pixmap.scaled(self.image_label.width(), self.image_label.height(),
                                         Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label.setPixmap(scaled_pixmap)
            
            # 在图像上绘制缺陷框
            result_image = image.copy()
            for detail in data['details']:
                x1, y1, x2, y2 = int(detail['bbox_x1']), int(detail['bbox_y1']), int(detail['bbox_x2']), int(detail['bbox_y2'])
                cv2.rectangle(result_image, (x1, y1), (x2, y2), (0, 255, 0), 2)
                label = f"{detail['class_name']}: {detail['confidence']:.2f}"
                cv2.putText(result_image, label, (x1, y1-5), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 1)
            
            # 再次转换并显示
            rgb_result = cv2.cvtColor(result_image, cv2.COLOR_BGR2RGB)
            qt_result = QtGui.QImage(rgb_result.data, w, h, bytes_per_line, QtGui.QImage.Format_RGB888)
            result_pixmap = QtGui.QPixmap.fromImage(qt_result)
            scaled_result = result_pixmap.scaled(self.image_label.width(), self.image_label.height(),
                                                  Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label.setPixmap(scaled_result)
        else:
            self.image_label.setText("无法加载图像")
        
        # 显示缺陷详情
        details = data['details']
        
        # 清空表格
        self.detail_table.setRowCount(0)
        
        # 填充表格
        for detail in details:
            row = self.detail_table.rowCount()
            self.detail_table.insertRow(row)
            
            self.detail_table.setItem(row, 0, QTableWidgetItem(str(detail['defect_index'])))
            self.detail_table.setItem(row, 1, QTableWidgetItem(f"{detail['confidence']:.3f}"))
            self.detail_table.setItem(row, 2, QTableWidgetItem(detail['class_name']))
            self.detail_table.setItem(row, 3, QTableWidgetItem(f"({detail['bbox_x1']:.1f}, {detail['bbox_y1']:.1f}, {detail['bbox_x2']:.1f}, {detail['bbox_y2']:.1f})"))
            self.detail_table.setItem(row, 4, QTableWidgetItem(f"{detail['center_x']:.1f}"))
            self.detail_table.setItem(row, 5, QTableWidgetItem(f"{detail['center_y']:.1f}"))
    
    def on_detail_item_clicked(self, item):
        """点击缺陷序号，以该缺陷为中心显示512*512窗口"""
        row = item.row()
        if row < 0:
            return
        
        # 获取缺陷信息
        defect_index = int(self.detail_table.item(row, 0).text())
        confidence = float(self.detail_table.item(row, 1).text())
        class_name = self.detail_table.item(row, 2).text()
        bbox_str = self.detail_table.item(row, 3).text()
        
        # 解析边界框坐标
        bbox_coords = bbox_str.strip('()').split(',')
        x1, y1, x2, y2 = float(bbox_coords[0].strip()), float(bbox_coords[1].strip()), \
                          float(bbox_coords[2].strip()), float(bbox_coords[3].strip())
        
        # 计算缺陷中心点
        center_x = (x1 + x2) / 2
        center_y = (y1 + y2) / 2
        
        # 检查是否有原始图像
        if self.current_image is None:
            QMessageBox.warning(self, "提示", "没有可用的原始图像")
            return
        
        # 获取图像尺寸
        img_h, img_w = self.current_image.shape[:2]
        
        # 计算512*512窗口的边界
        window_size = 512
        half_window = window_size // 2
        
        # 计算窗口起始点，确保不超出图像边界
        x1_crop = int(max(0, center_x - half_window))
        y1_crop = int(max(0, center_y - half_window))
        x2_crop = int(min(img_w, center_x + half_window))
        y2_crop = int(min(img_h, center_y + half_window))
        
        # 裁剪图像
        cropped_image = self.current_image[y1_crop:y2_crop, x1_crop:x2_crop].copy()
        
        # 计算实际裁剪区域的尺寸
        actual_h, actual_w = cropped_image.shape[:2]
        
        # 如果裁剪区域小于512*512，用灰色填充
        if actual_h < window_size or actual_w < window_size:
            padded_image = np.full((window_size, window_size, 3), 128, dtype=np.uint8)
            y_offset = (window_size - actual_h) // 2
            x_offset = (window_size - actual_w) // 2
            padded_image[y_offset:y_offset+actual_h, x_offset:x_offset+actual_w] = cropped_image
            cropped_image = padded_image
        
        # 在裁剪图像上绘制缺陷框（相对于裁剪区域的坐标）
        rel_x1 = int(x1 - x1_crop + (window_size - actual_w) // 2 if actual_w < window_size else int(x1 - x1_crop))
        rel_y1 = int(y1 - y1_crop + (window_size - actual_h) // 2 if actual_h < window_size else int(y1 - y1_crop))
        rel_x2 = rel_x1 + int(x2 - x1)
        rel_y2 = rel_y1 + int(y2 - y1)
        
        # 确保坐标在有效范围内
        rel_x1, rel_y1 = max(0, rel_x1), max(0, rel_y1)
        rel_x2, rel_y2 = min(window_size, rel_x2), min(window_size, rel_y2)
        
        # 绘制缺陷框（红色，更明显）
        cv2.rectangle(cropped_image, (rel_x1, rel_y1), (rel_x2, rel_y2), (0, 0, 255), 3)
        label = f"{class_name}: {confidence:.2f}"
        cv2.putText(cropped_image, label, (rel_x1, rel_y1-10), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 0, 255), 2)
        
        # 绘制中心点
        center_x_rel = int(center_x - x1_crop + (window_size - actual_w) // 2 if actual_w < window_size else int(center_x - x1_crop))
        center_y_rel = int(center_y - y1_crop + (window_size - actual_h) // 2 if actual_h < window_size else int(center_y - y1_crop))
        center_x_rel = max(0, min(window_size, center_x_rel))
        center_y_rel = max(0, min(window_size, center_y_rel))
        cv2.circle(cropped_image, (center_x_rel, center_y_rel), 5, (255, 0, 0), -1)
        
        # 绘制十字线
        cv2.line(cropped_image, (center_x_rel - 30, center_y_rel), (center_x_rel + 30, center_y_rel), (255, 0, 0), 2)
        cv2.line(cropped_image, (center_x_rel, center_y_rel - 30), (center_x_rel, center_y_rel + 30), (255, 0, 0), 2)
        
        # 绘制窗口边界
        cv2.rectangle(cropped_image, (0, 0), (window_size-1, window_size-1), (0, 255, 0), 2)
        
        # 转换颜色空间并显示
        rgb_image = cv2.cvtColor(cropped_image, cv2.COLOR_BGR2RGB)
        h, w, c = rgb_image.shape
        bytes_per_line = c * w
        qt_image = QtGui.QImage(rgb_image.data, w, h, bytes_per_line, QtGui.QImage.Format_RGB888)
        pixmap = QtGui.QPixmap.fromImage(qt_image)
        
        # 缩放图像以适应标签大小
        scaled_pixmap = pixmap.scaled(self.image_label.width(), self.image_label.height(),
                                     Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.image_label.setPixmap(scaled_pixmap)
        
        # 添加提示信息
        self.image_label.setToolTip(f"缺陷{defect_index} - {class_name}\n中心点: ({center_x:.1f}, {center_y:.1f})\n窗口: 512*512")

    def delete_selected_result(self):
        """删除选中的检测结果"""
        # 获取选中的行
        selected_rows = self.result_table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "提示", "请先选择要删除的检测结果")
            return
        
        # 确认删除
        reply = QMessageBox.question(self, "确认删除", f"确定要删除选中的{len(selected_rows)}条检测结果吗？\n此操作不可恢复。",
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        
        if reply == QMessageBox.No:
            return
        
        # 遍历选中的行并删除
        deleted_count = 0
        for row in selected_rows:
            result_id = int(self.result_table.item(row.row(), 0).text())
            success, msg = self.db_manager.delete_result(result_id)
            if success:
                deleted_count += 1
            else:
                QMessageBox.warning(self, "删除失败", f"删除结果ID {result_id} 失败：{msg}")
        
        if deleted_count > 0:
            QMessageBox.information(self, "删除成功", f"成功删除{deleted_count}条检测结果")
            # 刷新数据
            self.refresh_data()
            # 刷新可视化
            self.refresh_visualization()
            # 清空可视化区域
            self.image_label.setText("请选择一个检测结果查看可视化")
            self.current_image = None
            self.current_result_id = None

    def export_to_word(self):
        """将检测结果导出为Word文档（包含图表）"""
        # 获取所有检测结果
        results, msg = self.db_manager.get_all_results()
        if results is None or len(results) == 0:
            QMessageBox.warning(self, "提示", "没有可导出的检测结果！")
            return

        # 选择保存路径
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存Word文档",
            f"检测报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word文档 (*.docx)"
        )
        if not file_path:
            return

        try:
            # 创建Word文档
            doc = Document()

            # 设置中文字体
            from docx.oxml.ns import qn
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 设置页面边距（更宽以适应图表）
            section = doc.sections[0]
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

            # ========== 标题 ==========
            title = doc.add_heading('电路板漏铜检测报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 报告信息
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'用户：{self.username}')
            doc.add_paragraph()

            # ========== 计算统计数据 ==========
            total_images = len(results)
            total_defects = sum(r['total_defects'] for r in results)
            avg_defects = total_defects / total_images if total_images > 0 else 0

            # 计算良品率和收集数据
            pass_count = 0
            fail_count = 0
            defect_counts = []
            image_names = []
            defect_positions = []

            for result in results:
                details, _ = self.db_manager.get_defect_details(result['id'])
                defect_counts.append(result['total_defects'])
                # 截断过长的文件名
                short_name = result['image_name'][:15] + '...' if len(result['image_name']) > 15 else result[
                    'image_name']
                image_names.append(short_name)

                if details:
                    high_conf = [d for d in details if d['confidence'] > 0.8]
                    if len(high_conf) < 5:
                        pass_count += 1
                    else:
                        fail_count += 1
                    # 收集缺陷位置
                    for d in details:
                        defect_positions.append((d['center_x'], d['center_y']))
                else:
                    pass_count += 1

            yield_rate = (pass_count / total_images) * 100 if total_images > 0 else 0

            # ========== 1. 数据可视化图表（放在最前方） ==========
            doc.add_heading('一、数据可视化分析', level=1)

            # 临时保存图表到本地
            chart_dir = os.path.join(os.path.dirname(file_path), 'temp_charts')
            os.makedirs(chart_dir, exist_ok=True)

            # 1.1 缺陷数量柱状图
            doc.add_heading('1. 各图像缺陷数量统计', level=2)

            fig1 = Figure(figsize=(10, 5))
            ax1 = fig1.add_subplot(111)

            x = range(len(image_names))
            ax1.bar(x, defect_counts, color='steelblue', alpha=0.8)
            ax1.set_title('各图像缺陷数量分布')
            ax1.set_xlabel('图像')
            ax1.set_ylabel('缺陷数量')
            ax1.set_xticks(x)
            ax1.set_xticklabels(image_names, rotation=45, ha='right', fontsize=8)

            # 在柱子上显示数值
            for i, v in enumerate(defect_counts):
                ax1.text(i, v + 0.1, str(v), ha='center', va='bottom', fontsize=9)

            fig1.tight_layout()
            chart1_path = os.path.join(chart_dir, 'defect_count_chart.png')
            fig1.savefig(chart1_path, dpi=150, bbox_inches='tight')
            plt.close(fig1)

            doc.add_picture(chart1_path, width=Inches(6))
            doc.add_paragraph()

            # 1.2 良品率饼图
            doc.add_heading('2. 良品率统计', level=2)

            fig2 = Figure(figsize=(6, 5))
            ax2 = fig2.add_subplot(111)

            if pass_count + fail_count > 0:
                labels = ['合格品', '不合格品']
                sizes = [pass_count, fail_count]
                colors = ['#4CAF50', '#f44336']
                explode = (0.05, 0)

                ax2.pie(sizes, explode=explode, labels=labels, colors=colors,
                        autopct='%1.1f%%', shadow=True, startangle=90)
                ax2.axis('equal')
                ax2.set_title(f'良品率: {yield_rate:.1f}%')
            else:
                ax2.text(0.5, 0.5, '无数据', ha='center', va='center')

            fig2.tight_layout()
            chart2_path = os.path.join(chart_dir, 'yield_rate_chart.png')
            fig2.savefig(chart2_path, dpi=150, bbox_inches='tight')
            plt.close(fig2)

            doc.add_picture(chart2_path, width=Inches(4.5))
            doc.add_paragraph()

            # 1.3 缺陷位置分布散点图
            if defect_positions:
                doc.add_heading('3. 缺陷位置分布图', level=2)

                fig3 = Figure(figsize=(8, 6))
                ax3 = fig3.add_subplot(111)

                x_coords = [p[0] for p in defect_positions]
                y_coords = [p[1] for p in defect_positions]

                ax3.scatter(x_coords, y_coords, c='red', alpha=0.6, s=30)
                ax3.set_title('缺陷位置分布')
                ax3.set_xlabel('X坐标 (像素)')
                ax3.set_ylabel('Y坐标 (像素)')
                ax3.grid(True, alpha=0.3)

                fig3.tight_layout()
                chart3_path = os.path.join(chart_dir, 'defect_distribution_chart.png')
                fig3.savefig(chart3_path, dpi=150, bbox_inches='tight')
                plt.close(fig3)

                doc.add_picture(chart3_path, width=Inches(5.5))
                doc.add_paragraph()

            doc.add_page_break()

            # ========== 2. 统计摘要表格 ==========
            doc.add_heading('二、统计摘要', level=1)

            stats_table = doc.add_table(rows=5, cols=2)
            stats_table.style = 'Table Grid'
            stats_data = [
                ('总检测图像数', str(total_images)),
                ('总缺陷数', str(total_defects)),
                ('平均每图缺陷数', f'{avg_defects:.2f}'),
                ('合格品数', str(pass_count)),
                ('良品率', f'{yield_rate:.2f}%')
            ]
            for i, (label, value) in enumerate(stats_data):
                stats_table.rows[i].cells[0].text = label
                stats_table.rows[i].cells[1].text = value
                # 加粗左侧列
                for paragraph in stats_table.rows[i].cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            doc.add_paragraph()

            # ========== 3. 检测结果详情 ==========
            doc.add_heading('三、检测结果详情', level=1)

            for idx, result in enumerate(results, 1):
                doc.add_heading(f'图像 {idx}: {result["image_name"]}', level=2)

                # 获取该图像的缺陷详情
                details, _ = self.db_manager.get_defect_details(result['id'])

                # 基本信息表格
                info_data = [
                    ('缺陷总数', str(result['total_defects'])),
                    ('检测时间', str(result['detection_time'])),
                    ('窗口大小', result['window_size']),
                    ('输出路径', result['output_path'])
                ]

                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Table Grid'
                for i, (label, value) in enumerate(info_data):
                    info_table.rows[i].cells[0].text = label
                    info_table.rows[i].cells[1].text = value
                    for paragraph in info_table.rows[i].cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # 缺陷详情表格
                if details:
                    doc.add_paragraph('缺陷列表：')
                    defect_table = doc.add_table(rows=len(details) + 1, cols=6)
                    defect_table.style = 'Table Grid'

                    # 表头
                    headers = ['序号', '置信度', '类别', '边界框(x1,y1,x2,y2)', '中心点X', '中心点Y']
                    for i, header in enumerate(headers):
                        cell = defect_table.rows[0].cells[i]
                        cell.text = header
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # 填充数据
                    for i, detail in enumerate(details, 1):
                        defect_table.rows[i].cells[0].text = str(detail['defect_index'])
                        defect_table.rows[i].cells[1].text = f"{detail['confidence']:.3f}"
                        defect_table.rows[i].cells[2].text = detail['class_name']
                        defect_table.rows[i].cells[
                            3].text = f"({detail['bbox_x1']:.1f},{detail['bbox_y1']:.1f},{detail['bbox_x2']:.1f},{detail['bbox_y2']:.1f})"
                        defect_table.rows[i].cells[4].text = f"{detail['center_x']:.1f}"
                        defect_table.rows[i].cells[5].text = f"{detail['center_y']:.1f}"
                else:
                    doc.add_paragraph('无缺陷信息')

                doc.add_paragraph()  # 空行
                if idx < len(results):
                    doc.add_page_break()  # 每张图分页，最后一张不分页

            # ========== 4. 清理临时图表文件 ==========
            import shutil
            try:
                shutil.rmtree(chart_dir)
            except:
                pass

            # ========== 5. 保存文档 ==========
            doc.save(file_path)
            QMessageBox.information(self, "导出成功",
                                    f"报告已保存到：\n{file_path}\n\n报告包含：\n- 数据可视化图表\n- 统计摘要\n- 详细检测结果")

            # 询问是否打开
            reply = QMessageBox.question(self, "打开文件", "是否立即打开文档？",
                                         QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            if reply == QMessageBox.Yes:
                os.startfile(file_path)

        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "导出失败", f"导出Word文档时出错：\n{str(e)}")
    def refresh_visualization(self):
        """刷新数据可视化"""
        # 获取所有检测结果
        results, msg = self.db_manager.get_all_results()
        if results is None:
            QMessageBox.warning(self, "错误", f"获取数据失败：{msg}")
            return
        
        if not results:
            # 没有数据时显示空图表
            self.clear_all_canvas()
            self.update_stats(0, 0, 0, 0, 0, 0, 0)
            return
        
        # 计算统计数据
        total_images = len(results)
        total_defects = sum(result['total_defects'] for result in results)
        average_defects = total_defects / total_images if total_images > 0 else 0
        
        # 计算良品率（置信度大于0.8的缺陷数 < 5）
        pass_count = 0
        fail_count = 0
        high_conf_defects = []
        defect_positions = []
        
        for result in results:
            # 获取该图像的缺陷详情
            details, _ = self.db_manager.get_defect_details(result['id'])
            if details:
                # 统计置信度大于0.8的缺陷数
                high_conf = [d for d in details if d['confidence'] > 0.8]
                high_conf_defects.append(len(high_conf))
                
                # 统计缺陷位置
                for d in details:
                    defect_positions.append((d['center_x'], d['center_y']))
                
                # 判断是否合格
                if len(high_conf) < 5:
                    pass_count += 1
                else:
                    fail_count += 1
            else:
                pass_count += 1
        
        yield_rate = (pass_count / total_images) * 100 if total_images > 0 else 0
        
        # 更新统计信息
        self.update_stats(total_images, total_defects, average_defects, yield_rate, pass_count, fail_count, len(high_conf_defects))
        
        # 绘制缺陷数量柱状图
        self.plot_defect_count(results)
        
        # 绘制良品率饼图
        self.plot_yield_rate(pass_count, fail_count)
        
        # 绘制缺陷位置分布
        self.plot_defect_distribution(defect_positions)

    def update_stats(self, total_images, total_defects, average_defects, yield_rate, pass_count, fail_count, high_conf_count):
        """更新统计信息"""
        self.total_images_label.setText(f"总图像数: {total_images}")
        self.total_defects_label.setText(f"总缺陷数: {total_defects}")
        self.average_defects_label.setText(f"平均缺陷数: {average_defects:.2f}")
        self.yield_rate_label.setText(f"良品率: {yield_rate:.2f}%")
        self.pass_count_label.setText(f"合格品数: {pass_count}")
        self.fail_count_label.setText(f"不合格品数: {fail_count}")

    def clear_all_canvas(self):
        """清空所有图表"""
        # 清空缺陷数量柱状图
        self.defect_count_ax.clear()
        self.defect_count_ax.set_title('各图像缺陷数量')
        self.defect_count_ax.set_xlabel('图像')
        self.defect_count_ax.set_ylabel('缺陷数量')
        self.defect_count_ax.text(0.5, 0.5, '无数据', ha='center', va='center', transform=self.defect_count_ax.transAxes)
        self.defect_count_canvas.draw()
        
        # 清空良品率饼图
        self.yield_pie_ax.clear()
        self.yield_pie_ax.set_title('良品率')
        self.yield_pie_ax.text(0.5, 0.5, '无数据', ha='center', va='center', transform=self.yield_pie_ax.transAxes)
        self.yield_pie_canvas.draw()
        
        # 清空缺陷分布热力图
        self.defect_dist_ax.clear()
        self.defect_dist_ax.set_title('缺陷位置分布')
        self.defect_dist_ax.set_xlabel('X坐标')
        self.defect_dist_ax.set_ylabel('Y坐标')
        self.defect_dist_ax.text(0.5, 0.5, '无数据', ha='center', va='center', transform=self.defect_dist_ax.transAxes)
        self.defect_dist_canvas.draw()

    def plot_defect_count(self, results):
        """绘制缺陷数量折线图"""
        self.defect_count_ax.clear()
        
        # 准备数据
        image_names = [result['image_name'][:10] + '...' if len(result['image_name']) > 10 else result['image_name'] for result in results]
        defect_counts = [result['total_defects'] for result in results]
        x = range(len(image_names))
        
        # 绘制折线图
        self.defect_count_ax.plot(x, defect_counts, marker='o', markersize=8, linestyle='-', linewidth=2, color='b')
        self.defect_count_ax.set_title('各图像缺陷数量')
        self.defect_count_ax.set_xlabel('图像')
        self.defect_count_ax.set_ylabel('缺陷数量')
        
        # 设置X轴标签
        self.defect_count_ax.set_xticks(x)
        self.defect_count_ax.set_xticklabels(image_names, rotation=45, ha='right')
        
        # 调整y轴范围，上下增加空间
        if defect_counts:
            min_count = min(defect_counts)
            max_count = max(defect_counts)
            # 计算上下边距，增加边距比例使y轴更长
            margin = max(2, (max_count - min_count) * 0.4)  # 增加边距到40%
            y_min = max(0, min_count - margin)
            y_max = max_count + margin
            # 确保y轴至少有足够的空间
            if max_count < 5:
                y_max = max(5, y_max)  # 最小y轴最大值为5
            self.defect_count_ax.set_ylim(y_min, y_max)
        
        # 在每个数据点上添加标签
        for i, count in enumerate(defect_counts):
            self.defect_count_ax.text(i, count + 0.1, str(count), ha='center', va='bottom', fontsize=9, color='black')
        
        # 添加网格
        self.defect_count_ax.grid(alpha=0.3)
        
        # 调整布局
        self.defect_count_canvas.figure.tight_layout()
        self.defect_count_canvas.draw()

    def plot_yield_rate(self, pass_count, fail_count):
        """绘制良品率饼图"""
        self.yield_pie_ax.clear()
        
        if pass_count + fail_count == 0:
            self.yield_pie_ax.text(0.5, 0.5, '无数据', ha='center', va='center')
        else:
            labels = ['合格品', '不合格品']
            sizes = [pass_count, fail_count]
            colors = ['#4CAF50', '#f44336']
            explode = (0.1, 0)
            
            # 绘制饼图
            wedges, texts, autotexts = self.yield_pie_ax.pie(
                sizes, explode=explode, labels=labels, colors=colors,
                autopct='%1.1f%%', shadow=True, startangle=90
            )
            
            # 设置文本属性
            for text in texts:
                text.set_fontsize(12)
            for autotext in autotexts:
                autotext.set_fontsize(10)
                autotext.set_color('white')
            
            self.yield_pie_ax.axis('equal')  # 确保饼图是圆形
        
        self.yield_pie_ax.set_title('良品率')
        self.yield_pie_canvas.draw()

    def plot_defect_distribution(self, defect_positions):
        """绘制缺陷位置分布热力图"""
        self.defect_dist_ax.clear()
        
        if not defect_positions:
            self.defect_dist_ax.text(0.5, 0.5, '无数据', ha='center', va='center')
        else:
            # 提取X和Y坐标
            x = [pos[0] for pos in defect_positions]
            y = [pos[1] for pos in defect_positions]
            
            # 绘制散点图
            scatter = self.defect_dist_ax.scatter(x, y, c='red', alpha=0.6, s=50)
            
            # 添加颜色条
            self.defect_dist_ax.figure.colorbar(scatter, ax=self.defect_dist_ax, label='缺陷密度')
            
            # 添加网格
            self.defect_dist_ax.grid(alpha=0.3)
        
        self.defect_dist_ax.set_title('缺陷位置分布')
        self.defect_dist_ax.set_xlabel('X坐标')
        self.defect_dist_ax.set_ylabel('Y坐标')
        self.defect_dist_canvas.draw()

    def logout(self):
        """退出登录"""
        reply = QMessageBox.question(self, "确认退出", "确定要退出登录吗？",
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.close()
            # 重新显示登录窗口
            login_window = LoginWindow(self.db_manager)
            login_window.login_success.connect(lambda u: self.show_main_window(u))
            login_window.show()
    
    def show_main_window(self, username):
        """显示主窗口"""
        self.username = username
        self.user_label.setText(f"当前用户: {username}")
        self.setWindowTitle(f"电路板漏铜检测系统 - 当前用户: {username}")
        self.show()
        self.refresh_data()

if __name__ == "__main__":
    # 检查YOLOv8和PyTorch是否安装
    try:
        import ultralytics
        import torch
    except ImportError:
        print("请先安装依赖：pip install ultralytics torch pyqt5 PyQt5-sip")
        sys.exit(1)

    # 启动应用
    app = QApplication(sys.argv)
    
    # 创建数据库管理器
    db_manager = DatabaseManager()
    success, msg = db_manager.connect()
    if not success:
        QMessageBox.critical(None, "数据库连接失败", f"错误信息：{msg}\n程序将退出。")
        sys.exit(1)
    
    # 显示登录界面
    login_window = LoginWindow(db_manager)
    
    # 创建主窗口
    main_window = YoloV8Interface(db_manager, "")
    
    # 登录成功后显示主窗口
    def on_login_success(username):
        main_window.username = username
        main_window.user_label.setText(f"当前用户: {username}")
        main_window.setWindowTitle(f"电路板漏铜检测系统 - 当前用户: {username}")
        main_window.show()
    
    login_window.login_success.connect(on_login_success)
    
    # 显示登录界面
    login_window.show()
    
    sys.exit(app.exec_())